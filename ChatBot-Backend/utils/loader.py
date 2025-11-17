import os
from langchain_core.documents import Document
from PyPDF2 import PdfReader
import docx
import pandas as pd

def load_policy_files(data_folder: str):
    docs = []

    if not os.path.exists(data_folder):
        raise ValueError(f"Data folder '{data_folder}' not found.")

    files = [f for f in os.listdir(data_folder) if not f.startswith(".")]
    if not files:
        raise ValueError(f"No readable files found in '{data_folder}'.")

    for filename in files:
        path = os.path.join(data_folder, filename)
        ext = os.path.splitext(filename)[1].lower()

        try:
            if ext == ".pdf":
                text = ""
                reader = PdfReader(path)
                for page in reader.pages:
                    text += page.extract_text() or ""
                docs.append(Document(page_content=text, metadata={"source": filename}))

            elif ext in [".docx", ".doc"]:
                doc = docx.Document(path)
                text = "\n".join([p.text for p in doc.paragraphs])
                docs.append(Document(page_content=text, metadata={"source": filename}))

            elif ext in [".xlsx", ".xls"]:
                df = pd.read_excel(path)
                text = df.to_string(index=False)
                docs.append(Document(page_content=text, metadata={"source": filename}))

            elif ext in [".txt", ".md"]:
                with open(path, "r", encoding="utf-8") as f:
                    text = f.read()
                docs.append(Document(page_content=text, metadata={"source": filename}))

        except Exception as e:
            print(f"⚠️ Could not load {filename}: {e}")

    if not docs:
        raise ValueError(f"No supported files found in '{data_folder}'. Ensure you have PDF, DOCX, XLSX, or TXT files.")

    return docs
